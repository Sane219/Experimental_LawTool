"""
Text extraction service for the Legal Document Summarizer.
Handles extraction from PDF, DOCX, and TXT files with cleaning and validation.
"""

import re
from typing import Optional
from pathlib import Path

try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None


class TextExtractor:
    """Extracts and processes text from various document formats."""
    
    def __init__(self):
        """Initialize the text extractor."""
        self._legal_keywords = [
            'contract', 'agreement', 'party', 'parties', 'whereas', 'therefore',
            'shall', 'liability', 'obligation', 'clause', 'section', 'article',
            'defendant', 'plaintiff', 'court', 'judge', 'law', 'legal', 'statute',
            'regulation', 'compliance', 'breach', 'damages', 'remedy', 'jurisdiction',
            'governing', 'applicable', 'enforce', 'binding', 'valid', 'void',
            'terminate', 'termination', 'effective', 'date', 'execution'
        ]
    
    def extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using multiple methods for best results.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If PDF extraction fails
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        
        text = ""
        
        # Try pdfplumber first (better for complex layouts)
        if pdfplumber:
            try:
                text = self._extract_with_pdfplumber(file_path)
                if text.strip():
                    return self.clean_text(text)
            except Exception:
                pass  # Fall back to PyPDF2
        
        # Fall back to PyPDF2
        if PyPDF2:
            try:
                text = self._extract_with_pypdf2(file_path)
                if text.strip():
                    return self.clean_text(text)
            except Exception:
                pass
        
        # If both methods fail
        if not text.strip():
            raise Exception("Unable to extract text from PDF. The file may be corrupted or contain only images.")
        
        return self.clean_text(text)
    
    def extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file while preserving basic formatting context.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If DOCX extraction fails
        """
        if not Document:
            raise ImportError("python-docx library not available")
        
        if not Path(file_path).exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                raise Exception("No readable text found in DOCX file")
            
            full_text = "\n\n".join(text_parts)
            return self.clean_text(full_text)
            
        except Exception as e:
            if "No readable text found" in str(e):
                raise
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    
    def extract_from_txt(self, file_path: str) -> str:
        """
        Extract text from TXT file with encoding detection.
        
        Args:
            file_path: Path to the TXT file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If TXT extraction fails
        """
        if not Path(file_path).exists():
            raise FileNotFoundError(f"TXT file not found: {file_path}")
        
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                    if text.strip():
                        return self.clean_text(text)
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                raise Exception(f"Failed to read TXT file: {str(e)}")
        
        raise Exception("Unable to decode text file. Unsupported encoding.")
    
    def clean_text(self, raw_text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            raw_text: Raw extracted text
            
        Returns:
            Cleaned and normalized text
        """
        if not raw_text:
            return ""
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', raw_text)
        
        # Normalize line breaks first
        text = re.sub(r'\r\n|\r', '\n', text)
        
        # Remove excessive line breaks (more than 2 consecutive)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Clean up spaces within lines (but preserve line breaks)
        lines = []
        for line in text.split('\n'):
            # Remove excessive whitespace within each line
            cleaned_line = re.sub(r'\s+', ' ', line).strip()
            lines.append(cleaned_line)
        
        # Join lines back together
        text = '\n'.join(lines)
        
        # Remove empty lines at start and end
        text = text.strip()
        
        return text
    
    def validate_legal_content(self, text: str) -> bool:
        """
        Validate that the text appears to be legal document content.
        
        Args:
            text: Text to validate
            
        Returns:
            True if text appears to be legal content
        """
        if not text or len(text.strip()) < 50:
            return False
        
        # Convert to lowercase for keyword matching
        text_lower = text.lower()
        
        # Count legal keywords
        keyword_count = sum(1 for keyword in self._legal_keywords if keyword in text_lower)
        
        # Check for legal document patterns
        legal_patterns = [
            r'\b(whereas|therefore|shall|hereby)\b',
            r'\b(party|parties)\b.*\b(agree|contract)\b',
            r'\b(section|article|clause)\s+\d+',
            r'\b(effective|execution)\s+date\b',
            r'\b(governing|applicable)\s+law\b'
        ]
        
        pattern_matches = sum(1 for pattern in legal_patterns if re.search(pattern, text_lower))
        
        # Consider it legal content if:
        # - Has at least 3 legal keywords, OR
        # - Has at least 2 legal patterns, OR
        # - Has both keywords and patterns
        return keyword_count >= 3 or pattern_matches >= 2 or (keyword_count >= 1 and pattern_matches >= 1)
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from file based on its extension.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If extraction fails or format is unsupported
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            return self.extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_from_txt(file_path)
        else:
            raise Exception(f"Unsupported file format: {file_extension}")
    
    def _extract_with_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber library."""
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)
    
    def _extract_with_pypdf2(self, file_path: str) -> str:
        """Extract text using PyPDF2 library."""
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n\n".join(text_parts)